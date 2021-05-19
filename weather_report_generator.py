import requests
import pandas as pd
from bs4 import BeautifulSoup
import matplotlib.pyplot as plt
import re
from docx import Document
document = Document()

page=requests.get("https://mausam.imd.gov.in/")
c = page.content
soup=BeautifulSoup(c,"html.parser")
today=(soup.find(id='today'))
weather=(today.find_all(class_="capital"))


city=[item.find('h3').get_text() for item in weather]
temp=[item.find(class_='now').get_text() for item in weather]
wind=[item.find(class_='wind').get_text() for item in weather]
rain=[item.find(class_='minmax').get_text() for item in weather]


report=pd.DataFrame({
    'City':city,
    'Temperature':temp,
    'Wind':wind,
    'Rain':rain,
})

report.to_csv('report4.csv')

rain1=[]
for i in range(8):
    r=(int(re.search('[0-9]+',rain[i]).group(0)))
    rain1.insert(i,r)

temp1=[]
for i in range(8):
    r=(int(re.search('[0-9]+',temp[i]).group(0)))
    temp1.insert(i,r)

wind1=[]
for i in range(8):
    r=(int(re.search('[0-9]+',wind[i]).group(0)))
    wind1.insert(i,r)

plt.style.use("bmh")
plt.xlabel('City',fontsize=18)
plt.ylabel('Tempearture',fontsize=18)
plt.bar(city,temp1)
plt.savefig('Example.png')
plt.show()

plt.xlabel('Cities', fontsize=18)
plt.ylabel('Wind(km/hr)', fontsize=16)
plt.plot(city,wind1)
plt.savefig('Example1.png')
plt.show()

plt.pie(rain1, labels=city, radius=1.2,autopct='%0.01f%%', shadow=True)
plt.savefig('Example2.png')
plt.show()

document.add_picture('Example.png') 
document.add_picture('Example1.png') 
document.add_picture('Example2.png') 
document.save('Report.docx')
